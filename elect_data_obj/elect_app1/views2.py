# -*-coding:utf-8 -*-
import os

import docx
import docx2pdf
import pandas as pd
import pythoncom
from django.shortcuts import render
from django.http import HttpResponse, JsonResponse, FileResponse
from django.views import View
from django.db import connection
from docx.oxml.ns import qn
from docx.shared import Pt

from elect_app1.general import delete_extra_zero, region_all_data, sum_data, elect_start_data, fracture_data, \
    result_count_data, take_quantity
from elect_app1.models import RegionDistrict, ElectricSatrtSumData
from elect_app1.pdf_creata import left_start_png, plot_png, right_take_quantity


# Create your views here.


def get_overview_data(request):
    if request.method == 'GET':
        sql1 = """
        SELECT SUM( b.fire_generation ), SUM( b.water_generation ), SUM( b.take_generation ),
        SUM( b.wind_generation ), SUM( b.light_generation ) FROM
        (select * FROM region_district WHERE region = '{}') a
        INNER JOIN electric_satrt_sum_data b ON a.district = b.district where b.timing = '15:45:00'
        """

        sql2 = """
        SELECT  a.region, SUM( b.{} )
        FROM
        region_district a
        INNER JOIN electric_satrt_sum_data b ON a.district = b.district
        where b.timing = '15:45:00'
        GROUP BY a.region 
        """

        with connection.cursor() as cursor:  # with语句用于数据库操作
            cursor.execute(sql1.format('湖南省'))
            hunan_data = cursor.fetchall()
            cursor.execute(sql1.format('湖北省'))
            hubei_data = cursor.fetchall()
            cursor.execute(sql1.format('江西省'))
            jiangxi_data = cursor.fetchall()
            cursor.execute(sql1.format('河南省'))
            henan_data = cursor.fetchall()

            cursor.execute(sql2.format('generation_summation'))
            generation_summation_data = cursor.fetchall()
            cursor.execute(sql2.format('receive_generation'))
            receive_generation_data = cursor.fetchall()
            cursor.execute(sql2.format('real_time_generation'))
            real_time_generation_data = cursor.fetchall()
            cursor.execute(sql2.format('backup_generation_sum'))
            backup_generation_sum_data = cursor.fetchall()

        result_list = []
        for i, j in zip([hunan_data, hubei_data, jiangxi_data, henan_data],
                        ['hunan', 'hubei', 'jiangxi', 'henan']):
            region_dict = region_all_data(i, str(j))
            result_list.append(region_dict)

        for i, j in zip([generation_summation_data, receive_generation_data, real_time_generation_data,
                         backup_generation_sum_data],
                        ['generation_summation', 'receive_generation', 'real_time_generation',
                         'backup_generation_sum']):
            sum_dict = sum_data(i, j)
            result_list.append(sum_dict)

        print(result_list)
        result_data = {
            "message": '成功',
            "code": 200,
            'data': result_list,
        }
        return JsonResponse(result_data, json_dumps_params={'ensure_ascii': False},
                            content_type='application/json, charset=utf-8')


# 获取地区信息
def get_region(request):
    if request.method == 'GET':
        district_data = RegionDistrict.objects.all().values('region', 'district')
        area_list = []  # 转成list后json序列化
        for a in district_data:
            area_list.append({'region': a['region'], 'district': a['district']})
        # 然后通过jsonResponse返回给请求方, 这里是list而不是dict, 所以safe需要传入False.
        result_data = {
            "message": '成功',
            "code": 200,
            'data': area_list,
        }
        return JsonResponse(result_data, json_dumps_params={'ensure_ascii': False},
                            content_type='application/json, charset=utf-8')


# 获取初始数据 左1
def get_elect_start(request):
    if request.method == 'POST':
        region = request.POST.get('region')  # 省份
        district = request.POST.get('district')  # 片区
        datatimes = request.POST.get('datatimes')  # 日期

        df1, result_list = elect_start_data(district, region, datatimes)

        result_data = {
            "message": '成功',
            "code": 200,
            'data': result_list,
        }
        return JsonResponse(result_data, json_dumps_params={'ensure_ascii': False},
                            content_type='application/json, charset=utf-8')


# 断面数据
def get_fracture_data(request):
    flag = request.POST.get('flag')  # 标注 1 新能源2 最大供应
    datatimes = request.POST.get('datatimes')  # 日期

    df1, result_list = fracture_data(flag, datatimes)

    result_data = {
        "message": '成功',
        "code": 200,
        'data': result_list,
    }
    return JsonResponse(result_data, json_dumps_params={'ensure_ascii': False},
                        content_type='application/json, charset=utf-8')


# 结果数据
def get_result_data(request):
    if request.method == 'POST':
        region = request.POST.get('region')  # 省份
        district = request.POST.get('district')  # 片区
        datatimes = request.POST.get('datatimes')  # 日期
        flag = request.POST.get('flag')  # 标注 1 新能源2 最大供应
        df1, result_list = result_count_data(district, flag, region, datatimes)

        result_data = {
            "message": '成功',
            "code": 200,
            'data': result_list,
        }
        return JsonResponse(result_data, json_dumps_params={'ensure_ascii': False},
                            content_type='application/json, charset=utf-8')


# 获取抽蓄电量
def get_take_quantity(request):
    if request.method == 'POST':
        region = request.POST.get('region')  # 省份
        district = request.POST.get('district')  # 片区
        datatimes = request.POST.get('datatimes')  # 日期
        flag = request.POST.get('flag')  # 标注 1 新能源2 最大供应
        df, result_list = take_quantity(district, flag, region, datatimes)
        result_data = {
            "message": '成功',
            "code": 200,
            'data': result_list,
        }
        return JsonResponse(result_data, json_dumps_params={'ensure_ascii': False},
                            content_type='application/json, charset=utf-8')


def get_send_pdf(request):
    if request.method == 'GET':
        region = request.GET.get('region')  # 省份
        district = request.GET.get('district')  # 片区
        datatimes = request.GET.get('datatimes')  # 日期
        flag = request.GET.get('flag')  # 标注 1 新能源2 最大供应
        folder = os.getcwd()
        folder_path = folder + f'/pdf_data/{datatimes}'
        print(folder_path)
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        df1, result_list = elect_start_data(district, region, datatimes)
        left_start_png(df1, "generation", f'{folder_path}/left_data.png')
        plot_png(df1, "take_generation", "储能", f"{folder_path}/left_take.png")
        plot_png(df1, "receive_generation", "负荷", f"{folder_path}/left_load.png")

        df2, result_list = fracture_data(flag, datatimes)
        plot_png(df2, "fracture_data", "断面", f"{folder_path}/fracture.png")

        df3, result_list = result_count_data(district, flag, region, datatimes)
        left_start_png(df3, "quantity", f'{folder_path}/right_data.png')
        plot_png(df3, "take_quantity", "储能", f"{folder_path}/right_take_power.png")
        plot_png(df3, "load_power", "负荷", f"{folder_path}/right_load.png")

        df4, result_list = take_quantity(district, flag, region, datatimes)
        right_take_quantity(df4, f'{folder_path}/right_take_quantity.png')

        flag = int(flag)
        if flag == 1:
            mt11 = "新能源"
        else:
            mt11 = '最大供应'
        document = docx.Document()  # 创建一个Document对象（word文件）
        mt = f"""
        {region} 的 {district}片区 的{mt11} 的数据图表
        """
        document.add_paragraph(mt)
        try:
            document.add_picture(f'{folder_path}/left_data.png', width=docx.shared.Inches(6))
            document1 = document.add_paragraph('原始电源数据')
        except Exception as e:
            pass
        try:
            document.add_picture(f'{folder_path}/fracture.png', width=docx.shared.Inches(6))
            document1 = document.add_paragraph('原始电网数据')
        except Exception as e:
            pass
        try:
            document.add_picture(f'{folder_path}/left_take.png', width=docx.shared.Inches(6))
            document1 = document.add_paragraph('原始储能数据')
        except Exception as e:
            pass
        try:
            document.add_picture(f'{folder_path}/left_load.png', width=docx.shared.Inches(6))
            document1 = document.add_paragraph('原始负荷数据')
        except Exception as e:
            pass

        try:
            document.add_picture(f'{folder_path}/right_data.png', width=docx.shared.Inches(6))
            document1 = document.add_paragraph('计算后电源数据')
        except Exception as e:
            pass

        try:
            document.add_picture(f'{folder_path}/fracture.png', width=docx.shared.Inches(6))
            document1 = document.add_paragraph('计算后电网数据')
        except Exception as e:
            pass

        try:
            document.add_picture(f'{folder_path}/right_take_power.png', width=docx.shared.Inches(6))
            document1 = document.add_paragraph('计算后储能数据-发电功率')
        except Exception as e:
            pass

        try:
            document.add_picture(f'{folder_path}/right_take_quantity.png', width=docx.shared.Inches(6))
            document1 = document.add_paragraph('计算后储能数据-发电电量')
        except Exception as e:
            pass
        try:
            document.add_picture(f'{folder_path}/right_load.png', width=docx.shared.Inches(6))
            document1 = document.add_paragraph('计算后负荷数据')
        except Exception as e:
            pass

        # document.font.size = Pt(16)
        document.styles['Normal'].font.size = Pt(16)
        document.styles['Normal'].font.name = u'微软雅黑'
        document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        fn = f'{folder_path}/新型电力系统智能化调度图表.docx'
        file_path = f'{folder_path}/新型电力系统智能化调度图表.pdf'
        document.save(fn)

        pythoncom.CoInitialize()
        docx2pdf.convert(fn, file_path)

        file = open(file_path, 'rb')
        response = FileResponse(file)
        response['Content_type'] = "application/octet-stream"
        ##不加attachment是预览，不将文件下载下来
        response['Content-Disposition'] = f'filename="{file_path}"'
        # 加attachment是下载文件
        response['Content-Disposition'] = f'attachment;filename="{file_path}"'
        return response
